#!/usr/bin/env python3
"""
sync_benchmark_report.py - Synchronize benchmark results, compile Programming_Benchmark_Report.docx,
and facilitate repository updates and Git push.
"""

import os
import sys
import subprocess
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
REPORT_MD = os.path.join(REPO_ROOT, "Programming_Benchmark_Report.md")
REPORT_DOCX = os.path.join(REPO_ROOT, "Programming_Benchmark_Report.docx")
SUMMARY_PY = os.path.join(REPO_ROOT, "main_web_benchmark", "results", "generate_summary.py")

def run_summary_generator():
    """Step 1: Execute generate_summary.py to aggregate raw results."""
    if os.path.exists(SUMMARY_PY):
        print("[1/4] Running generate_summary.py...")
        subprocess.run([sys.executable, SUMMARY_PY], cwd=os.path.dirname(SUMMARY_PY), check=True)
    else:
        print("[1/4] generate_summary.py not found, skipping.")

def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def markdown_to_docx(md_path, docx_path):
    """Step 2: Convert Programming_Benchmark_Report.md into a Word Document (.docx)."""
    print(f"[2/4] Compiling {md_path} -> {docx_path}...")
    if not os.path.exists(md_path):
        print(f"Error: {md_path} does not exist.")
        return

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = docx.Document()

    # Set standard margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Style normal font
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'TH Sarabun New' if 'TH Sarabun New' in [f.name for f in doc.styles] else 'Calibri'
    normal_style.font.size = Pt(14)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    in_table = False
    table_data = []

    for line in lines:
        stripped = line.strip()

        # Handle Markdown Tables
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            # Skip separator lines like |---|---|
            if set(stripped.replace('|', '').replace('-', '').replace(':', '').strip()) == set():
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            table_data.append(cells)
            continue
        elif in_table:
            # End of table, write it to doc
            if table_data:
                cols_count = max(len(r) for r in table_data)
                tbl = doc.add_table(rows=len(table_data), cols=cols_count)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r_idx, row in enumerate(table_data):
                    for c_idx, cell_value in enumerate(row):
                        cell = tbl.cell(r_idx, c_idx)
                        cell.text = cell_value
                        # Header styling
                        if r_idx == 0:
                            set_cell_background(cell, "2B4C7E")
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        else:
                            if r_idx % 2 == 1:
                                set_cell_background(cell, "F4F6F9")
                doc.add_paragraph()  # spacer
            in_table = False
            table_data = []

        if not stripped:
            continue

        # Handle Headings
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:])
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:])
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2B, 0x4C, 0x7E)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:])
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x3B, 0x59, 0x98)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
        elif stripped.startswith('บทที่ '):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
        elif stripped.startswith('* ') or stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(stripped[2:])
            p.paragraph_format.space_after = Pt(2)
        elif stripped.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            p = doc.add_paragraph()
            p.add_run(stripped)
            p.paragraph_format.space_after = Pt(3)
        elif stripped.startswith('[') and ']' in stripped and stripped.endswith('.'):
            p = doc.add_paragraph()
            p.add_run(stripped)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(4)
        else:
            p = doc.add_paragraph()
            p.add_run(stripped)
            p.paragraph_format.space_after = Pt(4)

    # Flush any remaining table
    if in_table and table_data:
        cols_count = max(len(r) for r in table_data)
        tbl = doc.add_table(rows=len(table_data), cols=cols_count)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for r_idx, row in enumerate(table_data):
            for c_idx, cell_value in enumerate(row):
                cell = tbl.cell(r_idx, c_idx)
                cell.text = cell_value
                if r_idx == 0:
                    set_cell_background(cell, "2B4C7E")
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                elif r_idx % 2 == 1:
                    set_cell_background(cell, "F4F6F9")

    try:
        doc.save(docx_path)
        print(f"Successfully saved Word Document: {docx_path}")
    except PermissionError:
        alt_path = docx_path.replace(".docx", "_updated.docx")
        print(f"Warning: '{docx_path}' is locked (likely open in Microsoft Word).")
        print(f"Saving to alternative path: '{alt_path}'...")
        doc.save(alt_path)
        print(f"Successfully saved to: {alt_path}")

def git_commit_and_push(commit_msg="benchmarks: update data, sync report docx, and mirror README"):
    """Step 4: Stage, commit, and push changes to GitHub."""
    print("[3/4] Checking Git status...")
    status = subprocess.run(["git", "status", "-s"], capture_output=True, text=True, cwd=REPO_ROOT).stdout
    if not status.strip():
        print("No changes detected in working tree.")
        return

    print(f"[4/4] Committing and pushing to GitHub: '{commit_msg}'...")
    subprocess.run(["git", "add", "."], cwd=REPO_ROOT, check=True)
    res = subprocess.run(["git", "commit", "-m", commit_msg], cwd=REPO_ROOT, capture_output=True, text=True)
    print(res.stdout.strip())
    
    push_res = subprocess.run(["git", "push", "origin", "main"], cwd=REPO_ROOT, capture_output=True, text=True)
    if push_res.returncode == 0:
        print("Successfully pushed to origin/main!")
    else:
        print(f"Push output: {push_res.stderr.strip()}")

def main():
    import argparse
    parser = argparse.ArgumentParser(description="Synchronize benchmark data, update report .docx, and push to GitHub.")
    parser.add_argument("--summary-only", action="store_true", help="Only run generate_summary.py")
    parser.add_argument("--docx", action="store_true", help="Compile Programming_Benchmark_Report.docx")
    parser.add_argument("--push", action="store_true", help="Commit and push changes to GitHub")
    parser.add_argument("-m", "--message", default="benchmarks: update results, sync report docx, and mirror READMEs", help="Commit message")
    parser.add_argument("--all", action="store_true", default=True, help="Execute full sync pipeline")

    args = parser.parse_args()

    if args.summary_only:
        run_summary_generator()
    elif args.docx:
        markdown_to_docx(REPORT_MD, REPORT_DOCX)
    elif args.push:
        git_commit_and_push(args.message)
    else:
        run_summary_generator()
        markdown_to_docx(REPORT_MD, REPORT_DOCX)
        if args.push or "--push" in sys.argv:
            git_commit_and_push(args.message)

if __name__ == "__main__":
    main()
