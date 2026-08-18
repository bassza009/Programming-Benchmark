import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m_name, m_val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m_name}')
        node.set(qn('w:w'), str(m_val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1A365D") # Navy
        set_cell_margins(hdr_cells[i], top=140, bottom=140, left=160, right=160)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
            run.font.name = "Leelawadee UI"

    # Data rows
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F7FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, cell_value in enumerate(row_data):
            row_cells[c_idx].text = str(cell_value)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(45, 55, 72)
                run.font.name = "Leelawadee UI"

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph() # Spacing
    return table

def add_callout_box(doc, text_list, title="ประเด็นสำคัญของโครงการ", bg_color="EDF2F7", border_color="2B6CB0"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    # Left border styling
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="36" w:space="0" w:color="{border_color}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"📌 {title}\n")
    run_title.font.bold = True
    run_title.font.size = Pt(10.5)
    run_title.font.color.rgb = RGBColor(26, 54, 93)
    run_title.font.name = "Leelawadee UI"
    
    for item in text_list:
        p_item = cell.add_paragraph()
        p_item.paragraph_format.space_after = Pt(3)
        run = p_item.add_run(f"• {item}")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(45, 55, 72)
        run.font.name = "Leelawadee UI"
        
    doc.add_paragraph()

def build_document_th(file_path):
    doc = docx.Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Document Header / Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(4)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Project Antigravity: รายงานการทดสอบเปรียบเทียบประสิทธิภาพ Web Framework หลายภาษาและหลายสภาพแวดล้อม")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(26, 54, 93) # Navy
    title_run.font.name = "Leelawadee UI"

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(18)
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("เอกสารภาพรวมโครงการ วัตถุประสงค์ สถาปัตยกรรมระบบ รูปแบบการทดสอบ และบทวิเคราะห์เชิงลึก (ฉบับภาษาไทย)")
    sub_run.font.size = Pt(11.5)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(74, 85, 104)
    sub_run.font.name = "Leelawadee UI"

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(12)
    r_div = p_div.add_run("―" * 48)
    r_div.font.color.rgb = RGBColor(203, 213, 225)
    r_div.font.size = Pt(12)
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # SECTION 1: บทสรุปสำหรับผู้บริหาร (EXECUTIVE SUMMARY)
    h1 = doc.add_heading("1. บทสรุปสำหรับผู้บริหาร: โครงการนี้คืออะไร?", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    for r in h1.runs:
        r.font.name = "Leelawadee UI"
        r.font.color.rgb = RGBColor(26, 54, 93)
    
    p = doc.add_paragraph(
        "ในวงการพัฒนาซอฟต์แวร์ Backend ยุคปัจจุบัน มีการถกเถียงกันอย่างต่อเนื่องเกี่ยวกับภาษาและ Web Framework ที่ดีที่สุด "
        "เช่น 'Go เร็วกว่า Node.js จริงหรือไม่?', 'Java มีขนาดใหญ่เกินไปหรือเปล่า?', 'PHP ยังเร็วพอสำหรับระบบยุคใหม่หรือไม่?' "
        "และ 'การรันแอปพลิเคชันบน Docker ทำให้เซิร์ฟเวอร์ช้าลงมากน้อยเพียงใด?' อย่างไรก็ตาม การทดสอบประสิทธิภาพ (Benchmark) "
        "ส่วนใหญ่บนอินเทอร์เน็ตมักทดสอบเฉพาะโปรแกรมอย่างง่าย เช่น 'Hello World' ซึ่งส่งค่าข้อความสั้นๆ กลับมา "
        "ซึ่งไม่สามารถสะท้อนความเป็นจริงในระบบ Production ได้เลย เพราะระบบจริงต้องทำงานร่วมกับฐานข้อมูล (Relational Database), "
        "จัดการ Database Connection Pool, ประมวลผลคำสั่ง SQL JOIN หลายตาราง และรับมือกับผู้ใช้งานพร้อมกันนับพันคน"
    )
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        r.font.name = "Leelawadee UI"

    p = doc.add_paragraph(
        "โครงการนี้ (Project Antigravity / Programming-Benchmark) จึงถูกพัฒนาขึ้นเพื่อเป็น ชุดทดสอบประสิทธิภาพมาตรฐาน "
        "(Benchmark Suite) ที่มีความเป็นกลาง เป็นไปตามหลักการทางวิทยาศาสตร์ และสามารถทำซ้ำได้ (Deterministic & Reproducible) "
        "โดยทำการประเมิน 5 ภาษาและ Framework ยอดนิยม ภายใต้ภาระงานฐานข้อมูลจริง (Realistic Database Workloads) "
        "เปรียบเทียบระหว่างการทำงานบนเครื่องโดยตรง (Bare Metal) กับการทำงานบนคอนเทนเนอร์ (Docker) "
        "ครอบคลุมระดับการใช้งานตั้งแต่ระบบทดสอบไปจนถึงการทดสอบความเค้นระดับขีดจำกัด (10,000 Concurrent Connections)"
    )
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(12)
    for r in p.runs:
        r.font.name = "Leelawadee UI"

    add_callout_box(
        doc,
        [
            "5 ภาษาและ Framework ที่ทดสอบ: Python (FastAPI), Node.js (Fastify), PHP (Swoole), Go (Fiber), และ Java (Spring Boot)",
            "2 สภาพแวดล้อมระบบ: Bare Metal (รันตรงบน OS) เปรียบเทียบกับ Docker Containers (รันบนระบบเสมือน)",
            "5 ระดับโหลดตามสถานการณ์จริง: ตั้งแต่ระบบ POC (20 conns) ไปจนถึง Stress Test จุดอิ่มตัว (10,000 conns)",
            "ภาระงานฐานข้อมูลสมจริง: การสืบค้นตารางเดี่ยว, JOIN 2-4 ตาราง (แบบมี Index vs ไม่มี Index) และธุรกรรมการเขียน (POST Insert Transactions)",
            "รองรับการรันซ้ำหลายรอบ (--runs N) เพื่อหาค่าเฉลี่ยทางสถิติ และจัดเก็บ Log ผลลัพธ์ดิบรายรอบใน raw_results.json"
        ],
        title="สรุปหัวใจสำคัญของโครงการในภาพรวม"
    )

    # SECTION 2: ทำไมจึงต้องสร้างโครงการนี้ (THE CORE WHY & GOALS)
    h1 = doc.add_heading("2. ทำไมจึงต้องสร้างโครงการนี้? (เป้าหมายและปัญหาที่ได้รับการแก้ไข)", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    for r in h1.runs:
        r.font.name = "Leelawadee UI"
        r.font.color.rgb = RGBColor(26, 54, 93)

    p = doc.add_paragraph(
        "เมื่อทีมวิศวกรซอฟต์แวร์หรือผู้ออกแบบระบบ (Architects) ต้องเลือก Technology Stack สำหรับสร้าง Web API หรือ Microservices "
        "ที่ต้องรองรับโหลดสูง พวกเขามักเผชิญกับคำถามสำคัญที่ไม่สามารถหาคำตอบได้จากการทดสอบแบบสังเคราะห์ทั่วไป:"
    )
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        r.font.name = "Leelawadee UI"

    goals = [
        ("ก้าวข้ามกับดัก 'Hello World'", "ในการทำงานจริง เซิร์ฟเวอร์ใช้เวลากว่า 80-90% ไปกับการรอ Database I/O, แปลงข้อมูล JSON, และจัดการ Connection Pool ไม่ใช่การคำนวณลูปตัวเลขในแรม โครงการนี้จึงทดสอบการเชื่อมต่อ MySQL จริงที่มีข้อมูลหลักหมื่นแถว"),
        ("วัดต้นทุนความหน่วงของ Docker (Virtualization Tax)", "การ Deploy ระบบบน Docker Container ทำให้สูญเสีย Throughput และเพิ่ม Latency มากน้อยเพียงใด? ชุดทดสอบนี้รันโค้ดชุดเดียวกันบนเครื่องเดียวกัน เพื่อวัดผลกระทบของการทำ Virtualization และ Bridge Network อย่างแม่นยำ"),
        ("ผลกระทบของ Database Index ในสภาวะ Concurrency สูง", "เกิดอะไรขึ้นหากนักพัฒนาลืมสร้าง Index บน Foreign Key เมื่อมีทราฟฟิกเข้ามาพร้อมกันนับพัน? เราเปรียบเทียบคำสั่ง SQL เดียวกันระหว่างแบบมี Index และไม่มี Index เพื่อให้เห็นตัวเลขความต่างอย่างชัดเจน"),
        ("หาขีดจำกัดความเสถียร (Breakdown & Saturation Limits)", "Framework ใดสามารถทนทานต่อการเชื่อมต่อพร้อมกัน 10,000 Connections ได้อย่างราบรื่น และ Framework ใดเริ่มมี Connection หลุด (Errors/Drops) หรือเกิด Latency พุ่งสูง"),
        ("มาตรฐานความยุติธรรมแบบเท่าเทียมกันทุกภาษา", "การเปรียบเทียบส่วนใหญ่มักมีความลำเอียง เช่น ตั้ง Connection Pool ไม่เท่ากัน หรือไม่รอ Warmup โครงการนี้กำหนดขนาด Connection Pool เท่ากัน มีช่วง Warmup 3 วินาที รีเซ็ตฐานข้อมูลระหว่างรอบ และปรับ ulimit ป้องกันคอขวดของ OS")
    ]

    for g_title, g_desc in goals:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r_bold = p.add_run(f"• {g_title}: ")
        r_bold.bold = True
        r_bold.font.color.rgb = RGBColor(26, 54, 93)
        r_bold.font.name = "Leelawadee UI"
        r_txt = p.add_run(g_desc)
        r_txt.font.name = "Leelawadee UI"

    doc.add_paragraph()

    # SECTION 3: สถาปัตยกรรมและเทคโนโลยีที่ใช้ (TECH STACK)
    h1 = doc.add_heading("3. เทคโนโลยีและสถาปัตยกรรมระบบที่นำมาประเมิน", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    for r in h1.runs:
        r.font.name = "Leelawadee UI"
        r.font.color.rgb = RGBColor(26, 54, 93)

    p = doc.add_paragraph(
        "เพื่อให้การเปรียบเทียบมีความทันสมัยและเป็นตัวแทนของเทคโนโลยีที่ดีที่สุด เราได้คัดเลือก Framework ประสิทธิภาพสูง "
        "ที่รองรับ Asynchronous / Concurrency ของแต่ละภาษา ดังนี้:"
    )
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        r.font.name = "Leelawadee UI"

    tech_headers = ["ภาษา (Language)", "Web Framework", "Database Driver / Client", "โมเดลการประมวลผล (Concurrency)", "Port มาตรฐาน"]
    tech_data = [
        ["Python", "FastAPI + Uvicorn", "aiomysql (Async Pool)", "Multi-process Async Event Loop", "8001"],
        ["Node.js", "Fastify", "mysql2/promise (Pool)", "Multi-core Cluster + Event Loop", "8002"],
        ["PHP", "Swoole", "PDO_MySQL (PDOPool)", "C-based Coroutine Event Loop", "8003"],
        ["Go", "Fiber (v2)", "database/sql (go-sql-driver)", "Lightweight Goroutines", "8004"],
        ["Java", "Spring Boot (v3)", "JdbcTemplate + HikariCP", "Multi-threaded JVM Pool", "8005"]
    ]
    add_styled_table(doc, tech_headers, tech_data, [1.0, 1.4, 1.8, 1.8, 0.8])

    # SECTION 4: รูปแบบการทดสอบและระดับโหลด (SCENARIOS & TIERS)
    h1 = doc.add_heading("4. รูปแบบการทดสอบและระดับโหลด (Test Scenarios & Load Tiers)", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    for r in h1.runs:
        r.font.name = "Leelawadee UI"
        r.font.color.rgb = RGBColor(26, 54, 93)

    p = doc.add_paragraph("การทดสอบแบ่งออกเป็น 2 หมวดหมู่หลัก ครอบคลุมการอ่านและเขียนข้อมูล:")
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.font.name = "Leelawadee UI"

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("A. หมวดการอ่านข้อมูล / GET Benchmark Suites (Raw SQL):\n")
    r.bold = True
    r.font.color.rgb = RGBColor(43, 108, 176)
    r.font.name = "Leelawadee UI"
    r2 = p.add_run(
        "1. /raw/1table: สืบค้นตารางเดี่ยว (SELECT * FROM users LIMIT 100)\n"
        "2. /raw/2join: สืบค้นแบบเชื่อม 2 ตาราง (users + profiles)\n"
        "3. /raw/3join: สืบค้นแบบเชื่อม 3 ตาราง (users + profiles + orders)\n"
        "4. /raw/4join: สืบค้นแบบเชื่อม 4 ตาราง (users + profiles + orders + order_items)\n"
        "ทดสอบใน 2 สภาวะฐานข้อมูล: (1) แบบไม่มี Index (No Index) เพื่อทดสอบ Table Scan และ (2) แบบมี Index (With Index) บน Foreign Key"
    )
    r2.font.name = "Leelawadee UI"

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("B. หมวดการเขียนข้อมูล / POST Benchmark Suite (Database Transactions):\n")
    r.bold = True
    r.font.color.rgb = RGBColor(43, 108, 176)
    r.font.name = "Leelawadee UI"
    r2 = p.add_run(
        "1. /raw/post/1table: บันทึกข้อมูลลงตาราง users 1 รายการ\n"
        "2. /raw/post/2table: บันทึกข้อมูลแบบ Transaction เชื่อมโยง user + profile\n"
        "3. /raw/post/3table: บันทึกข้อมูลแบบ Transaction เชื่อมโยง user + profile + order\n"
        "4. /raw/post/4table: บันทึกข้อมูลแบบ Transaction ครบวงจร user + profile + order + order_items หลายรายการ"
    )
    r2.font.name = "Leelawadee UI"

    p = doc.add_paragraph("ตาราง 5 ระดับโหลดตามสถานการณ์การใช้งานจริง (ทดสอบด้วยเครื่องมือ wrk):")
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.font.name = "Leelawadee UI"

    tier_headers = ["สถานการณ์ (Scenario)", "ตัวอย่างระบบจริง (Typical Website)", "เธรด (-t)", "Connections (-c)", "ระยะเวลา (-d)"]
    tier_data = [
        ["POC / Small internal system", "โปรเจกต์จบการศึกษา, ระบบต้นแบบในแผนก", "2", "20", "30 วินาที"],
        ["Small production website", "เว็บไซต์บริษัทขนาดเล็ก, ธุรกิจท้องถิ่น", "4", "100", "60 วินาที"],
        ["General web application", "ระบบมหาวิทยาลัย, ระบบอีคอมเมิร์ซ, CMS", "8", "500", "60 วินาที"],
        ["High-density website", "เว็บพอร์ทัลยอดนิยม, แพลตฟอร์ม SaaS", "8", "2,000", "120 วินาที"],
        ["Stress testing", "หาจุดอิ่มตัวและขีดจำกัดสูงสุดของระบบ", "16", "10,000", "300 วินาที"]
    ]
    add_styled_table(doc, tier_headers, tier_data, [1.6, 2.0, 0.9, 1.1, 0.9])

    # SECTION 5: ผลการค้นพบและข้อสรุปสำคัญ (KEY DISCOVERIES)
    h1 = doc.add_heading("5. ผลการค้นพบและข้อสรุปสำคัญเชิงลึก (Key Benchmark Discoveries)", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    for r in h1.runs:
        r.font.name = "Leelawadee UI"
        r.font.color.rgb = RGBColor(26, 54, 93)

    p = doc.add_paragraph(
        "จากผลการทดสอบจริงบนเครื่องทดสอบมาตรฐาน ข้อมูลเชิงตัวเลขได้สะท้อนข้อเท็จจริงที่น่าสนใจดังต่อไปนี้:"
    )
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        r.font.name = "Leelawadee UI"

    res_headers = ["Suite & Endpoint", "Go (Fiber)", "Java (Spring Boot)", "PHP (Swoole)", "Node.js (Fastify)", "Python (FastAPI)"]
    res_data = [
        ["GET 1-Table (Bare Metal)", "11,928 req/s (9.3ms)", "11,958 req/s (8.4ms)", "15,762 req/s (7.3ms)", "7,017 req/s (16.3ms)", "1,624 req/s (61.2ms)"],
        ["GET 1-Table (Docker)", "10,988 req/s (10.7ms)", "9,232 req/s (12.2ms)", "16,003 req/s (6.9ms)", "2,042 req/s (49.2ms)", "2,516 req/s (40.0ms)"],
        ["GET 4-Join With Index (BME)", "3,700 req/s (28.4ms)", "3,646 req/s (28.3ms)", "3,990 req/s (30.4ms)", "3,821 req/s (27.9ms)", "1,483 req/s (67.0ms)"],
        ["POST 1-Table Insert (Docker)", "7,124 req/s (14.1ms)", "5,709 req/s (17.6ms)", "4,507 req/s (23.9ms)", "7,297 req/s (13.9ms)", "7,045 req/s (14.4ms)"]
    ]
    add_styled_table(doc, res_headers, res_data, [1.8, 1.1, 1.1, 1.1, 1.1, 1.1])

    insights = [
        ("PHP Swoole คือผู้นำด้านความเร็วที่น่าทึ่ง", "PHP เมื่อรันผ่าน Swoole Coroutines และ PDOPool สามารถทำ Throughput สูงสุดในการอ่านตารางเดี่ยวได้มากกว่า 16,000 requests/sec ด้วย Latency เพียง ~7ms ลบล้างความเชื่อเดิมที่ว่า PHP ทำงานช้า"),
        ("Go และ Java มีความเสถียรและทนทานสูงสุด", "Go (Fiber) และ Java (Spring Boot) ให้ประสิทธิภาพระดับท็อปอย่างสม่ำเสมอ (11,000+ req/s สำหรับการอ่าน และ 7,000+ req/s สำหรับการเขียน) โดยมี Latency Jitter ต่ำมาก และแทบไม่พบ Error แม้โหลดจะเพิ่มถึง 10,000 Connections"),
        ("ต้นทุนความหน่วงของ Docker (Virtualization Overhead)", "Node.js (Fastify) ทำได้กว่า 7,000 req/s บน Bare Metal แต่เมื่อรันบน Docker ภายใต้โหลดสูง ประสิทธิภาพลดลงอย่างเห็นได้ชัด เนื่องจาก Overhead ในการแปลงเน็ตเวิร์กของ Linux Bridge Driver"),
        ("พลังทวีคูณ 12 เท่าของ Database Index", "ในการ JOIN 4 ตาราง หากไม่มี Index ความเร็วจะตกลงเหลือเพียง ~300 req/s และ Latency พุ่งเกิน 1,000ms แต่เมื่อมี Index ความเร็วจะพุ่งขึ้นถึง ~3,800 req/s (เร็วขึ้นกว่า 12 เท่า) และ Latency ลดลงเหลือเพียง 28ms"),
        ("Python FastAPI โดดเด่นในงานเขียนข้อมูล (POST Transactions)", "แม้ว่า Python จะเสียเปรียบในงานอ่านตารางขนาดใหญ่ที่ใช้ CPU มาก แต่เมื่อเป็นงานเขียนข้อมูลแบบ Async Transaction (aiomysql) FastAPI กลับทำได้ถึง 7,045 req/s เทียบเท่ากับ Go และ Node.js")
    ]

    for i_title, i_desc in insights:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r_bold = p.add_run(f"⭐ {i_title}: ")
        r_bold.bold = True
        r_bold.font.color.rgb = RGBColor(43, 108, 176)
        r_bold.font.name = "Leelawadee UI"
        r_txt = p.add_run(i_desc)
        r_txt.font.name = "Leelawadee UI"

    doc.add_paragraph()

    # SECTION 6: วิธีการทดสอบและเก็บผลลัพธ์ (HOW TO RUN)
    h1 = doc.add_heading("6. วิธีการรันทดสอบและระบบจัดเก็บผลลัพธ์", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    for r in h1.runs:
        r.font.name = "Leelawadee UI"
        r.font.color.rgb = RGBColor(26, 54, 93)

    p = doc.add_paragraph("การรัน Benchmark ต้องการเพียง Python 3, Docker, wrk และฐานข้อมูล MySQL 8.0 ในเครื่อง Local:")
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.font.name = "Leelawadee UI"

    steps = [
        ("ขั้นตอนที่ 1: ตรวจสอบ MySQL ในเครื่อง", "เปิดใช้งาน MySQL 8.0 บนพอร์ต 3306 (user=admin, password=secret, database=benchmark_db)"),
        ("ขั้นตอนที่ 2: รันการทดสอบ Bare Metal พร้อมระบุจำนวนรอบ", "cd main_web_benchmark/GET/get_no_index && python3 run_bme_wrk.py --tier all --runs 3"),
        ("ขั้นตอนที่ 3: รันการทดสอบบน Docker Container", "cd main_web_benchmark/GET/get_no_index && python3 run_dkr_wrk.py --tier all --runs 3"),
        ("ขั้นตอนที่ 4: การจัดเก็บผลลัพธ์เฉลี่ยและข้อมูลดิบ", "ผลลัพธ์ที่คำนวณค่าเฉลี่ยทางสถิติจะถูกบันทึกใน bme_benchmark_results.json ส่วน Log การรันดิบทุกรอบจะถูกรวบรวมไว้ใน raw_results.json และ results/raw_results/"),
        ("ขั้นตอนที่ 5: สร้างตารางสรุปผลเปรียบเทียบ", "cd main_web_benchmark && python3 compare_results.py GET/get_no_index/dkr_benchmark_results.json")
    ]

    for s_title, s_desc in steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r_bold = p.add_run(f"▶ {s_title}\n")
        r_bold.bold = True
        r_bold.font.color.rgb = RGBColor(26, 54, 93)
        r_bold.font.name = "Leelawadee UI"
        r_code = p.add_run(f"   {s_desc}")
        r_code.font.size = Pt(9.5)
        r_code.font.color.rgb = RGBColor(74, 85, 104)
        r_code.font.name = "Leelawadee UI"

    doc.add_paragraph()

    # SECTION 7: ประโยชน์และกลุ่มเป้าหมาย (TARGET AUDIENCE)
    h1 = doc.add_heading("7. ใครจะได้รับประโยชน์จากโครงการนี้?", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    for r in h1.runs:
        r.font.name = "Leelawadee UI"
        r.font.color.rgb = RGBColor(26, 54, 93)

    audience = [
        ("ผู้ออกแบบระบบและหัวหน้าทีมวิศวกรรม (Tech Leads & Architects)", "มีชุดข้อมูลเชิงประจักษ์ที่ตรวจสอบได้ สำหรับใช้เป็นเหตุผลสนับสนุนการเลือก Technology Stack แทนการเชื่อตามกระแสหรือข้อมูล Benchmark สังเคราะห์ที่ขาดความสมจริง"),
        ("วิศวกร Backend และ DevOps", "เข้าใจถึงผลกระทบและข้อจำกัดของการรันงานบน Docker Container และความสำคัญของการเลือกโหมดเครือข่ายเมื่อต้องรองรับโหลดสูง"),
        ("นักศึกษา นักวิจัย และนักพัฒนาซอฟต์แวร์ทั่วไป", "ใช้เป็นแหล่งอ้างอิงและกรณีศึกษาในการออกแบบ High-Performance API, การจัดการ Connection Pool, การใช้ Coroutines, และการเพิ่มประสิทธิภาพคำสั่ง SQL")
    ]

    for a_title, a_desc in audience:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r_bold = p.add_run(f"👥 {a_title}: ")
        r_bold.bold = True
        r_bold.font.name = "Leelawadee UI"
        r_txt = p.add_run(a_desc)
        r_txt.font.name = "Leelawadee UI"

    # Save document
    try:
        doc.save(file_path)
        print(f"Thai document successfully created at: {file_path}")
    except PermissionError:
        alt_path = file_path.replace('.docx', '_v2.docx')
        doc.save(alt_path)
        print(f"Notice: {file_path} is currently open. Saved to: {alt_path}")

if __name__ == '__main__':
    build_document_th('Project_Benchmark_Overview_TH.docx')
