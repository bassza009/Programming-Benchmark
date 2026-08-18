import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m_name, m_val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m_name}')
        node.set(qn('w:w'), str(m_val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1A365D") # Navy
        set_cell_margins(hdr_cells[i], top=140, bottom=140, left=160, right=160)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
            run.font.name = "Calibri"

    # Data rows
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F7FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, cell_value in enumerate(row_data):
            row_cells[c_idx].text = str(cell_value)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(45, 55, 72)
                run.font.name = "Calibri"

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph() # Spacing
    return table

def add_callout_box(doc, text_list, title="KEY TAKEAWAY", bg_color="EDF2F7", border_color="2B6CB0"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    # Left border styling
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="36" w:space="0" w:color="{border_color}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"📌 {title}\n")
    run_title.font.bold = True
    run_title.font.size = Pt(10.5)
    run_title.font.color.rgb = RGBColor(26, 54, 93)
    
    for item in text_list:
        p_item = cell.add_paragraph()
        p_item.paragraph_format.space_after = Pt(3)
        run = p_item.add_run(f"• {item}")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(45, 55, 72)
        
    doc.add_paragraph()

def build_document(file_path):
    doc = docx.Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Document Header / Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Project Antigravity: Multi-Language & Multi-Environment Web Benchmark")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(26, 54, 93) # Navy
    title_run.font.name = "Calibri"

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(20)
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Comprehensive Project Overview, Methodology, and Performance Insights Explained in Plain English")
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(74, 85, 104)
    sub_run.font.name = "Calibri"

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(14)
    r_div = p_div.add_run("―" * 48)
    r_div.font.color.rgb = RGBColor(203, 213, 225)
    r_div.font.size = Pt(12)
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # SECTION 1: EXECUTIVE SUMMARY
    h1 = doc.add_heading("1. Executive Summary: What is this project?", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "Modern backend development offers dozens of programming languages and frameworks, each claiming to be the fastest, "
        "most scalable, and most efficient. However, most public benchmarks test simple 'Hello World' endpoints in isolation, "
        "which completely fails to reflect real-world production realities where servers interact with relational databases, "
        "manage connection pools, process complex multi-table joins, and handle thousands of concurrent requests."
    )
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph(
        "This project, titled Project Antigravity (Programming-Benchmark), is an open-source, deterministic, and scientifically "
        "rigorous performance benchmark suite. It measures the true throughput, latency, concurrency limits, and resource efficiency "
        "of 5 major programming languages and web frameworks under realistic database-driven workloads across both Bare Metal (host OS) "
        "and Docker Containerized environments."
    )
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(12)

    add_callout_box(
        doc,
        [
            "5 Programming Languages Tested: Python (FastAPI), Node.js (Fastify), PHP (Swoole), Go (Fiber), and Java (Spring Boot).",
            "2 Operating Environments: Bare Metal (Direct OS) vs. Docker Containerized (Virtual/Isolated).",
            "3 Real-World Load Tiers: Light (100 conns), High-Load (1,000 conns), and Extreme Stress (10,000 conns).",
            "Realistic Workloads: Single-table reads, 2 to 4-table SQL JOINs, Indexed vs. Non-Indexed queries, and Multi-table Write/Insert Transactions."
        ],
        title="CORE PROJECT SUMMARY AT A GLANCE"
    )

    # SECTION 2: THE CORE "WHY" & GOALS
    h1 = doc.add_heading("2. Why Did We Build This? (The Core Goals & Problems Solved)", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph(
        "When software architects and engineering teams choose a technology stack for high-throughput web APIs or microservices, "
        "they face several critical questions that standard benchmarks fail to answer:"
    )
    p.paragraph_format.space_after = Pt(8)

    goals = [
        ("The 'Hello World' Trap", "In real production, servers spend 80-90% of their time waiting for database I/O, serializing JSON, and managing connection pools—not calculating synthetic loops. This suite tests real MySQL database I/O with realistic datasets (10,000+ rows)."),
        ("The Docker Virtualization Tax", "How much performance do we actually lose when containerizing our applications in Docker vs running on bare metal? This project tests both environments side-by-side using identical hardware and configurations to measure the exact containerization overhead."),
        ("Database Index Impact Under Extreme Load", "What happens when developers forget to add secondary indexes on foreign keys? We directly compare queries with and without database indexes across 100 to 10,000 concurrent connections."),
        ("Concurrency Breakdown Limits", "Where do frameworks break? Which frameworks handle 10,000 simultaneous connections gracefully, and which ones drop connections or crash with socket errors?"),
        ("Fairness & Standardization", "Most comparisons are flawed because connection pool sizes, worker counts, or JIT warmups differ. This project standardizes connection pools, enforces process isolation, implements 3-second warmup phases, and resets databases between tests.")
    ]

    for g_title, g_desc in goals:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r_bold = p.add_run(f"• {g_title}: ")
        r_bold.bold = True
        r_bold.font.color.rgb = RGBColor(26, 54, 93)
        p.add_run(g_desc)

    doc.add_paragraph()

    # SECTION 3: THE TECHNOLOGY STACK & ARCHITECTURE
    h1 = doc.add_heading("3. Evaluated Technologies & System Architecture", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph(
        "To ensure a fair and direct comparison, we selected the leading high-performance asynchronous or concurrent web framework "
        "for each of the top 5 backend languages:"
    )
    p.paragraph_format.space_after = Pt(8)

    tech_headers = ["Language", "Framework", "Database Driver / Client", "Concurrency Model", "Default Port"]
    tech_data = [
        ["Python", "FastAPI + Uvicorn", "aiomysql (Async Pool)", "Multi-process Async Event Loop", "8001"],
        ["Node.js", "Fastify", "mysql2/promise (Pool)", "Multi-core Cluster + Event Loop", "8002"],
        ["PHP", "Swoole", "PDO_MySQL (PDOPool)", "C-based Coroutine Event Loop", "8003"],
        ["Go", "Fiber (v2)", "database/sql (go-sql-driver)", "Lightweight Goroutines", "8004"],
        ["Java", "Spring Boot (v3)", "JdbcTemplate + HikariCP", "Multi-threaded JVM Pool", "8005"]
    ]
    add_styled_table(doc, tech_headers, tech_data, [1.0, 1.4, 1.8, 1.8, 0.8])

    # SECTION 4: TEST SCENARIOS & LOAD TIERS
    h1 = doc.add_heading("4. Test Scenarios and Load Testing Tiers", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph("The benchmark evaluates two primary types of operations across multiple difficulty levels:")
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("A. Read / GET Benchmark Suites (Raw SQL):\n")
    r.bold = True
    r.font.color.rgb = RGBColor(43, 108, 176)
    p.add_run(
        "1. /raw/1table: Single-table query (SELECT * FROM users LIMIT 100)\n"
        "2. /raw/2join: 2-table JOIN query (users + profiles)\n"
        "3. /raw/3join: 3-table JOIN query (users + profiles + orders)\n"
        "4. /raw/4join: 4-table JOIN query (users + profiles + orders + order_items)\n"
        "Tested in two database states: (1) No Indexes (table scan stress) and (2) With Indexes (optimized B-Tree lookups)."
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("B. Write / POST Benchmark Suite (Database Transactions):\n")
    r.bold = True
    r.font.color.rgb = RGBColor(43, 108, 176)
    p.add_run(
        "1. /raw/post/1table: Single-table insert into users table.\n"
        "2. /raw/post/2table: Transactional insert creating a user and linked profile record.\n"
        "3. /raw/post/3table: Transactional insert creating user + profile + order.\n"
        "4. /raw/post/4table: Transactional insert creating user + profile + order + multiple order items."
    )

    p = doc.add_paragraph("All suites are evaluated across 5 standardized production load scenarios using the wrk load generator:")
    p.paragraph_format.space_after = Pt(6)

    tier_headers = ["Scenario", "Typical Website", "Threads (-t)", "Connections (-c)", "Duration (-d)"]
    tier_data = [
        ["POC / Small internal system", "Thesis project, department website prototype", "2", "20", "30s"],
        ["Small production website", "Small company local business", "4", "100", "60s"],
        ["General web application", "University system e-commerce CMS", "8", "500", "60s"],
        ["High-density website", "Popular portals SaaS platforms", "8", "2,000", "120s"],
        ["Stress testing", "Find saturation point", "16", "10,000", "300s"]
    ]
    add_styled_table(doc, tier_headers, tier_data, [1.6, 2.0, 0.9, 1.1, 0.9])

    # SECTION 5: KEY BENCHMARK RESULTS & DISCOVERIES
    h1 = doc.add_heading("5. Key Benchmark Results & Discoveries in Plain English", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph(
        "Based on extensive empirical testing on dedicated benchmark infrastructure, here is what the data reveals:"
    )
    p.paragraph_format.space_after = Pt(8)

    res_headers = ["Suite & Endpoint", "Go (Fiber)", "Java (Spring Boot)", "PHP (Swoole)", "Node.js (Fastify)", "Python (FastAPI)"]
    res_data = [
        ["GET 1-Table (Bare Metal)", "11,928 req/s (9.3ms)", "11,958 req/s (8.4ms)", "15,762 req/s (7.3ms)", "7,017 req/s (16.3ms)", "1,624 req/s (61.2ms)"],
        ["GET 1-Table (Docker)", "10,988 req/s (10.7ms)", "9,232 req/s (12.2ms)", "16,003 req/s (6.9ms)", "2,042 req/s (49.2ms)", "2,516 req/s (40.0ms)"],
        ["GET 4-Join With Index (BME)", "3,700 req/s (28.4ms)", "3,646 req/s (28.3ms)", "3,990 req/s (30.4ms)", "3,821 req/s (27.9ms)", "1,483 req/s (67.0ms)"],
        ["POST 1-Table Insert (Docker)", "7,124 req/s (14.1ms)", "5,709 req/s (17.6ms)", "4,507 req/s (23.9ms)", "7,297 req/s (13.9ms)", "7,045 req/s (14.4ms)"]
    ]
    add_styled_table(doc, res_headers, res_data, [1.8, 1.1, 1.1, 1.1, 1.1, 1.1])

    insights = [
        ("PHP Swoole is a Hidden Speed Demon", "PHP with Swoole coroutines and PDOPool achieved the highest raw throughput on single-table reads (over 16,000 req/s with ~7ms latency), completely defying traditional beliefs that PHP is slow."),
        ("Go and Java Provide Unshakable Consistency", "Go (Fiber) and Java (Spring Boot) consistently maintained elite performance (11,000+ req/s on reads and 7,000+ req/s on writes) with minimal latency jitter and near-zero connection errors."),
        ("The Virtualization Tax on Node.js", "Node.js (Fastify) ran at 7,000+ req/s on Bare Metal, but dropped significantly inside Docker containers under high thread counts due to Linux bridge network packet translation overhead."),
        ("Database Indexes Are Essential Under Concurrency", "In 3-table and 4-table joins without indexes, query throughput dropped from 3,800 req/s down to ~300 req/s, and average latency jumped from 28ms to 1,000ms+. Secondary indexes on join keys provided a 12x performance boost!"),
        ("Python FastAPI is Optimized for Writes/Async", "While Python lagged in CPU-bound heavy table scans, FastAPI with aiomysql performed exceptionally well in transactional write workloads (POST 1-table: 7,045 req/s), matching Go and Node.js.")
    ]

    for i_title, i_desc in insights:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r_bold = p.add_run(f"⭐ {i_title}: ")
        r_bold.bold = True
        r_bold.font.color.rgb = RGBColor(43, 108, 176)
        p.add_run(i_desc)

    doc.add_paragraph()

    # SECTION 6: HOW TO RUN
    h1 = doc.add_heading("6. How to Run the Benchmarks Yourself", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph("Running the benchmarks requires only Python 3, Docker, wrk, and a local MySQL 8.0 instance:")
    p.paragraph_format.space_after = Pt(6)

    steps = [
        ("Step 1: Ensure Local MySQL is Running", "Start a local MySQL 8.0 instance on port 3306 with credentials: user=admin, password=secret, database=benchmark_db."),
        ("Step 2: Run Bare Metal Benchmarks", "cd main_web_benchmark/GET/get_no_index && python3 run_bme_wrk.py --tier all --runs 3"),
        ("Step 3: Run Docker Benchmarks", "cd main_web_benchmark/GET/get_no_index && python3 run_dkr_wrk.py --tier all --runs 3"),
        ("Step 4: View Averaged & Raw Results", "Averaged metrics are in bme_benchmark_results.json; full raw iteration data is saved in raw_results.json (and results/raw_results/)."),
        ("Step 5: Generate Comparison Tables", "cd main_web_benchmark && python3 compare_results.py GET/get_no_index/dkr_benchmark_results.json")
    ]

    for s_title, s_desc in steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r_bold = p.add_run(f"▶ {s_title}\n")
        r_bold.bold = True
        r_bold.font.color.rgb = RGBColor(26, 54, 93)
        r_code = p.add_run(f"   {s_desc}")
        r_code.font.size = Pt(9.5)
        r_code.font.color.rgb = RGBColor(74, 85, 104)

    doc.add_paragraph()

    # SECTION 7: TARGET AUDIENCE & VALUE
    h1 = doc.add_heading("7. Who Benefits From This Project?", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)

    audience = [
        ("Software Architects & Tech Leads", "Provides empirical, reproducible data to justify technology stack choices rather than relying on internet hype or superficial synthetic benchmarks."),
        ("Backend & DevOps Engineers", "Demonstrates the exact performance trade-offs of Docker containerization and network modes (--network host vs bridge) under heavy socket loads."),
        ("Students & Computer Scientists", "Serves as an educational reference architecture for high-performance API design, asynchronous I/O, coroutines, and connection pool optimization.")
    ]

    for a_title, a_desc in audience:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r_bold = p.add_run(f"👥 {a_title}: ")
        r_bold.bold = True
        p.add_run(a_desc)

    # Save document
    try:
        doc.save(file_path)
        print(f"Document successfully created at: {file_path}")
    except PermissionError:
        alt_path = file_path.replace('.docx', '_new.docx')
        doc.save(alt_path)
        print(f"Notice: {file_path} is currently open in Word. Generated updated document at: {alt_path}")

if __name__ == '__main__':
    build_document('Project_Benchmark_Overview.docx')
